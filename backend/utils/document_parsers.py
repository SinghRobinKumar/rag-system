"""
Document parsers for various file formats.
Each parser extracts plain text from a specific file type.
"""
import os
import json
import csv
import io
from pathlib import Path
from typing import Optional

from backend.config import SUPPORTED_EXTENSIONS


def parse_document(file_path: str) -> Optional[str]:
    """
    Parse a document and return its text content.
    Routes to the appropriate parser based on file extension.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return None

    try:
        if ext == ".pdf":
            return _parse_pdf(file_path)
        elif ext in (".txt", ".md"):
            return _parse_text(file_path)
        elif ext == ".csv":
            return _parse_csv(file_path)
        elif ext in (".docx", ".doc"):
            return _parse_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return _parse_xlsx(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            return _parse_image(file_path)
        elif ext in (".html", ".htm"):
            return _parse_html(file_path)
        elif ext == ".json":
            return _parse_json(file_path)
        else:
            return None
    except Exception as e:
        print(f"[Parser Error] Failed to parse {file_path}: {e}")
        return None


def _parse_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF with enhanced table detection.
    Tables are converted to markdown format for better LLM understanding.
    """
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc):
            page_parts = [f"[Page {page_num + 1}]"]

            # Try to detect and extract tables first
            try:
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for table in tabs.tables:
                        df = table.extract()
                        if df and len(df) > 0:
                            # Build markdown table
                            md_table = _table_to_markdown(df)
                            if md_table:
                                page_parts.append(md_table)
            except Exception:
                pass  # Fall through to plain text if table detection fails

            # Also get full page text for non-table content
            page_text = page.get_text("text")
            if page_text.strip():
                # Clean up excessive whitespace/newlines from PDF extraction
                lines = [line.strip() for line in page_text.split("\n") if line.strip()]
                page_parts.append("\n".join(lines))

            if len(page_parts) > 1:  # More than just the page header
                text_parts.append("\n".join(page_parts))

    return "\n\n".join(text_parts)


def _table_to_markdown(table_data: list[list]) -> str:
    """Convert a 2D table (list of rows) to a markdown table string."""
    if not table_data or len(table_data) < 1:
        return ""

    # Clean cells: replace None with empty string, strip whitespace
    cleaned = []
    for row in table_data:
        cleaned_row = [str(cell).strip() if cell else "" for cell in row]
        # Skip completely empty rows
        if any(cell for cell in cleaned_row):
            cleaned.append(cleaned_row)

    if not cleaned:
        return ""

    # Normalize column count
    max_cols = max(len(row) for row in cleaned)
    for row in cleaned:
        while len(row) < max_cols:
            row.append("")

    # Build markdown table
    header = cleaned[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * max_cols) + " |",
    ]
    for row in cleaned[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _parse_text(file_path: str) -> str:
    """Read plain text / markdown files."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _parse_csv(file_path: str) -> str:
    """Convert CSV to readable text format."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return ""

    # Use first row as headers
    headers = rows[0]
    text_parts = [f"CSV File: {Path(file_path).name}", f"Columns: {', '.join(headers)}"]

    for i, row in enumerate(rows[1:], 1):
        row_text = " | ".join(
            f"{headers[j] if j < len(headers) else f'Col{j}'}: {val}"
            for j, val in enumerate(row)
        )
        text_parts.append(f"Row {i}: {row_text}")

    return "\n".join(text_parts)


def _parse_docx(file_path: str) -> str:
    """Extract text from Word documents."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also try to extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _parse_xlsx(file_path: str) -> str:
    """Extract text from Excel files."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"Sheet: {sheet_name}")

        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            str_row = [str(cell) if cell is not None else "" for cell in row]
            if i == 0:
                headers = str_row
                text_parts.append(f"Columns: {', '.join(headers)}")
            else:
                row_text = " | ".join(
                    f"{headers[j] if j < len(headers) else f'Col{j}'}: {val}"
                    for j, val in enumerate(str_row)
                    if val
                )
                if row_text:
                    text_parts.append(f"Row {i}: {row_text}")

    wb.close()
    return "\n".join(text_parts)


def _parse_image(file_path: str) -> str:
    """
    Extract text from images using OCR (Tesseract).
    Falls back gracefully if tesseract is not installed.
    """
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text if text.strip() else f"[Image: {Path(file_path).name} - no text detected]"
    except ImportError:
        return f"[Image: {Path(file_path).name} - OCR not available, install pytesseract]"
    except Exception as e:
        return f"[Image: {Path(file_path).name} - OCR failed: {e}]"


def _parse_html(file_path: str) -> str:
    """Extract text from HTML files (basic tag stripping)."""
    import re

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()

    # Remove script and style blocks
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
    html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL)
    # Remove tags
    text = re.sub(r"<[^>]+>", " ", html)
    # Clean whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _parse_json(file_path: str) -> str:
    """Convert JSON to readable text."""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    return json.dumps(data, indent=2, ensure_ascii=False)
